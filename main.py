import json
import re
from io import BytesIO
from docx import Document
import fitz
from json_helper import InputData as input


def extract_text_from_file(file):
    """Extract text from a file (PDF or DOCX), including tables."""
    try:
        # Determine file extension
        if isinstance(file, str):  # If file is a path
            file_extension = file.split(".")[-1].lower()
        else:  # If file is an in-memory object (e.g., from Streamlit)
            file_extension = file.name.split(".")[-1].lower()

        print(f"File extension: {file_extension}")  # Debug: Print file extension

        if file_extension == "pdf":
            return extract_text_from_pdf(file)
        elif file_extension == "docx":
            return extract_text_from_docx(file)
        else:
            raise ValueError("Unsupported file format. Only PDF and DOCX files are supported.")
    except Exception as e:
        print(f"Error in extract_text_from_file: {e}")  # Debug: Print any errors
        return ""


def extract_text_from_pdf(file):
    """Extract text from a PDF file."""
    try:
        if isinstance(file, str):
            doc = fitz.open(file)
        else:
            doc = fitz.open(stream=file.read(), filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text()
        print(f"Extracted PDF text: {text}")  # Debug: Print extracted text
        return text
    except Exception as e:
        print(f"Error in extract_text_from_pdf: {e}")
        return ""


def extract_text_from_docx(file):
    """Extract text from a DOCX file."""
    try:
        if isinstance(file, str):
            doc = Document(file)
        else:
            file.seek(0)  # Reset file pointer to the beginning
            doc = Document(BytesIO(file.read()))

        text = ""

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"

        print(f"Extracted DOCX text: {text}")  # Debug: Print extracted text
        return text
    except Exception as e:
        print(f"Error in extract_text_from_docx: {e}")
        return ""


def preprocess_text(text: str) -> str:
    """Clean text while preserving structure and convert it into a single long string."""
    # Remove all newline characters and extra spaces
    text = re.sub(r'\s+', ' ', text)  # Replace multiple spaces with a single space
    text = text.replace('\n', '')  # Remove all newline characters
    return text.strip()


def validate_and_clean_json(data: str) -> dict:
    """Improved JSON cleaning with field validation"""
    try:
        # Extract JSON using more precise pattern
        json_match = re.search(r'\{[^{}]*\{.*?\}[^{}]*\}|\{.*?\}', data, re.DOTALL)
        if not json_match:
            raise ValueError("No JSON content found in the data.")

        json_str = json_match.group()

        # Fix common LLM output issues
        json_str = (json_str
                    .replace('“', '"')
                    .replace('”', '"')
                    .replace("‘", "'")
                    .replace("’", "'")
                    .replace('\n', ' ')
                    )

        # Validate required fields exist
        parsed = json.loads(json_str)

        # Ensure "education" is a list
        if "education" in parsed and not isinstance(parsed["education"], list):
            parsed["education"] = [parsed["education"]]

        return parsed

    except Exception as e:
        print(f"Error cleaning and validating JSON: {e}")
        return None


def run_resume_extractor(file_path):
    text = extract_text_from_file(file_path)
    processed_text = preprocess_text(text)

    llm = input.llm()

    # Split the processed text into chunks of 4200 characters
    chunks = [processed_text[i:i + 4200] for i in range(0, len(processed_text), 4200)]

    all_data = []

    for idx, chunk in enumerate(chunks):
        data = llm.invoke(input.input_data(chunk))

        # Convert the data to a string (if it's not already)
        data_str = str(data)

        # Extract and clean the JSON content
        cleaned_json = validate_and_clean_json(data_str)

        if cleaned_json:
            all_data.append(cleaned_json)

    if all_data:
        merged_output = all_data[0]

        for data in all_data[1:]:
            merged_output["professional_experience"].extend(data["professional_experience"])

        with open('merged_output.json', 'w') as json_file:
            json.dump(merged_output, json_file, indent=4)


run_resume_extractor(r"C:\Users\PC\Downloads\resume example Ahmed Al.docx")

